"""
Resume-specific service layer.

Handles file validation and delegates to the LLM service for analysis.
"""

from __future__ import annotations

from services.llm_service import generate_resume_feedback
from services.mock_service import mock_resume_history


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


import io

def validate_resume_file(filename: str, size: int) -> str | None:
    """Return an error message if the file is invalid, else None."""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return f"Unsupported file type '{ext}'. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
    if size > MAX_FILE_SIZE:
        return f"File too large ({size / 1024 / 1024:.1f} MB). Maximum is {MAX_FILE_SIZE / 1024 / 1024:.0f} MB."
    return None


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        import pypdf
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"[resume_service] Error extracting PDF text: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import docx
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        print(f"[resume_service] Error extracting DOCX text: {e}")
        return ""


async def analyze_resume(file_name: str, file_bytes: bytes, target_role: str = "Software Engineer") -> dict:
    """Validate and analyze an uploaded resume."""
    error = validate_resume_file(file_name, len(file_bytes))
    if error:
        raise ValueError(error)

    ext = "." + file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    
    file_text = ""
    if ext == ".pdf":
        file_text = extract_text_from_pdf(file_bytes)
    elif ext in {".docx", ".doc"}:
        file_text = extract_text_from_docx(file_bytes)
    else:
        try:
            file_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            file_text = ""

    return await generate_resume_feedback(file_name, file_text, target_role)


def get_resume_history() -> list[dict]:
    """Return resume analysis history (mock data for now)."""
    return mock_resume_history()

